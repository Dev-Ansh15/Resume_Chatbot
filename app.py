import uuid
import os
import json
import oracledb
from flask import Flask, render_template, request, jsonify, session, send_from_directory
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from reportlab.lib.units import inch

# --- Configuration ---
app = Flask(__name__)
app.secret_key = 'your_very_secret_key_here'  # Required for session management
app.config['UPLOAD_FOLDER'] = 'generated_resumes'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# --- Oracle Database Connection ---
# !! Update with your Oracle DB credentials !!
ORACLE_USER = "resume_app"
ORACLE_PASSWORD = "Steve2011"
# DSN (Data Source Name) format: "hostname:port/service_name"
# For Oracle XE, the service_name is often the PDB name, e.g., "XEPDB1"
ORACLE_DSN = "localhost:1521/XEPDB1"

# Create a connection pool (recommended for web applications)
try:
    pool = oracledb.create_pool(user=ORACLE_USER, password=ORACLE_PASSWORD, dsn=ORACLE_DSN, min=2, max=5, increment=1)
except oracledb.DatabaseError as e:
    print("Error creating Oracle connection pool:", e)
    pool = None

# --- Conversation Module ---
# Defines the questions and the order (state machine)
CONVERSATION_FLOW = [
    {'state': 'ask_name', 'prompt': 'Welcome! To start, what is your full name?', 'key': 'full_name'},
    {'state': 'ask_email', 'prompt': 'Got it. What is your email address?', 'key': 'email'},
    {'state': 'ask_phone', 'prompt': 'What is your phone number?', 'key': 'phone_number'},
    {'state': 'ask_linkedin', 'prompt': 'What is your LinkedIn profile URL? (Type "skip" if none)', 'key': 'linkedin_profile'},
    {'state': 'ask_summary', 'prompt': 'Great. Now, write a brief professional summary about yourself.', 'key': 'summary'},
    {'state': 'ask_education', 'prompt': 'Let\'s add your education. What is your most recent school/university?', 'key': 'education'},
    {'state': 'ask_degree', 'prompt': 'What was your degree or field of study?', 'key': 'education_degree'},
    {'state': 'ask_edu_dates', 'prompt': 'When did you graduate? (e.g., May 2024)', 'key': 'education_dates'},
    {'state': 'ask_more_education', 'prompt': 'Do you want to add another education entry? (yes/no)', 'key': 'education_loop'},
    {'state': 'ask_experience', 'prompt': 'Now for work experience. What was your most recent job title?', 'key': 'experience'},
    {'state': 'ask_company', 'prompt': 'What was the company name?', 'key': 'experience_company'},
    {'state': 'ask_exp_dates', 'prompt': 'What were the dates? (e.g., Jan 2022 - Present)', 'key': 'experience_dates'},
    {'state': 'ask_exp_desc', 'prompt': 'Describe one or two key responsibilities or achievements.', 'key': 'experience_desc'},
    {'state': 'ask_more_experience', 'prompt': 'Do you want to add another work experience? (yes/no)', 'key': 'experience_loop'},
    {'state': 'ask_skills', 'prompt': 'Finally, list your top skills, separated by commas (e.g., Python, Project Management, SQL)', 'key': 'skills'},
    {'state': 'complete', 'prompt': 'All done! Generating your resume...', 'key': None},
]

def get_current_state_info(state_name):
    """Finds the prompt and data key for the current state."""
    for item in CONVERSATION_FLOW:
        if item['state'] == state_name:
            return item
    return CONVERSATION_FLOW[0] # Default to start

# --- Database Module (Oracle) ---
def save_resume_to_oracle(resume_data):
    """Saves the completed resume data to the Oracle database."""
    if not pool:
        print("Database pool is not initialized.")
        return None

    sql = """
    INSERT INTO resumes (
        session_id, full_name, email, phone_number, linkedin_profile, 
        summary, education, experience, skills
    ) VALUES (
        :1, :2, :3, :4, :5, :6, :7, :8, :9
    ) RETURNING resume_id INTO :10
    """
    
    # Prepare data for insertion
    # Convert lists of dicts (education, experience) to JSON strings for CLOB storage
    education_json = json.dumps(resume_data.get('education_list', []))
    experience_json = json.dumps(resume_data.get('experience_list', []))
    skills_json = json.dumps(resume_data.get('skills', '').split(','))
    
    try:
        with pool.acquire() as connection:
            with connection.cursor() as cursor:
                # Create a variable to hold the returned resume_id
                new_id_var = cursor.var(oracledb.NUMBER)
                
                cursor.execute(sql, [
                    session['session_id'],
                    resume_data.get('full_name'),
                    resume_data.get('email'),
                    resume_data.get('phone_number'),
                    resume_data.get('linkedin_profile'),
                    resume_data.get('summary'),
                    education_json,
                    experience_json,
                    skills_json,
                    new_id_var  # Bind variable for the RETURNING clause
                ])
                connection.commit()
                
                # Get the new ID
                new_resume_id = new_id_var.getvalue()[0]
                print(f"Successfully saved resume with ID: {new_resume_id}")
                return new_resume_id
    except oracledb.DatabaseError as e:
        print(f"Error inserting data into Oracle: {e}")
        return None

# --- Resume Generator Module ---
def generate_docx(data, filename):
    """Generates a .docx resume."""
    doc = Document()
    doc.add_heading(data.get('full_name', 'Your Name'), 0)
    
    contact_info = f"{data.get('email', '')} | {data.get('phone_number', '')}"
    if data.get('linkedin_profile', 'skip') != 'skip':
        contact_info += f" | {data.get('linkedin_profile', '')}"
    doc.add_paragraph(contact_info)

    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(data.get('summary', ''))

    doc.add_heading('Work Experience', level=1)
    for exp in data.get('experience_list', []):
        doc.add_heading(f"{exp['title']} at {exp['company']}", level=2)
        p = doc.add_paragraph()
        p.add_run(exp['dates']).style = 'Emphasis'
        doc.add_paragraph(exp['desc'])

    doc.add_heading('Education', level=1)
    for edu in data.get('education_list', []):
        doc.add_heading(edu['school'], level=2)
        doc.add_paragraph(f"{edu['degree']} ({edu['dates']})")

    doc.add_heading('Skills', level=1)
    skills = ", ".join([s.strip() for s in data.get('skills', '').split(',')])
    doc.add_paragraph(skills)
    
    doc.save(filename)

def generate_pdf(data, filename):
    """Generates a .pdf resume."""
    doc = SimpleDocTemplate(filename)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(data.get('full_name', 'Your Name'), styles['h1']))
    contact_info = f"{data.get('email', '')} | {data.get('phone_number', '')}"
    if data.get('linkedin_profile', 'skip') != 'skip':
        contact_info += f" | {data.get('linkedin_profile', '')}"
    story.append(Paragraph(contact_info, styles['BodyText']))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph('Professional Summary', styles['h2']))
    story.append(Paragraph(data.get('summary', ''), styles['BodyText']))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph('Work Experience', styles['h2']))
    for exp in data.get('experience_list', []):
        story.append(Paragraph(f"{exp['title']} at {exp['company']}", styles['h3']))
        story.append(Paragraph(exp['dates'], styles['Italic']))
        story.append(Paragraph(exp['desc'], styles['BodyText']))
        story.append(Spacer(1, 0.1 * inch))

    story.append(Paragraph('Education', styles['h2']))
    for edu in data.get('education_list', []):
        story.append(Paragraph(edu['school'], styles['h3']))
        story.append(Paragraph(f"{edu['degree']} ({edu['dates']})", styles['BodyText']))
        story.append(Spacer(1, 0.1 * inch))

    story.append(Paragraph('Skills', styles['h2']))
    skills = ", ".join([s.strip() for s in data.get('skills', '').split(',')])
    story.append(Paragraph(skills, styles['BodyText']))

    doc.build(story)

# --- Controller (Flask Routes) ---

@app.route("/")
def index():
    """Serves the main chat page."""
    # Initialize the session for the conversation
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
    session['state'] = CONVERSATION_FLOW[0]['state']
    session['resume_data'] = {}
    session['temp_education'] = {}
    session['temp_experience'] = {}
    return render_template("index.html")

@app.route("/chat", methods=["POST"])
def chat():
    """Main chat API endpoint."""
    user_message = request.json['message'].strip()
    current_state = session.get('state', CONVERSATION_FLOW[0]['state'])
    state_info = get_current_state_info(current_state)
    
    # --- State Machine Logic ---
    next_state_name = ""
    bot_response = ""

    # Check for 'skip'
    if user_message.lower() == 'skip' and state_info['key'] not in ['full_name', 'email']:
        user_message = 'skip'
    
    # Save data from the previous state
    if state_info['key']:
        session['resume_data'][state_info['key']] = user_message

    # --- Education Loop ---
    if current_state == 'ask_education':
        session['temp_education'] = {'school': user_message}
        next_state_name = 'ask_degree'
    elif current_state == 'ask_degree':
        session['temp_education']['degree'] = user_message
        next_state_name = 'ask_edu_dates'
    elif current_state == 'ask_edu_dates':
        session['temp_education']['dates'] = user_message
        if 'education_list' not in session['resume_data']:
            session['resume_data']['education_list'] = []
        session['resume_data']['education_list'].append(session['temp_education'])
        session['temp_education'] = {}
        next_state_name = 'ask_more_education'
    elif current_state == 'ask_more_education':
        if user_message.lower() == 'yes':
            next_state_name = 'ask_education'
        else:
            next_state_name = 'ask_experience' # Move to next major section

    # --- Experience Loop ---
    elif current_state == 'ask_experience':
        session['temp_experience'] = {'title': user_message}
        next_state_name = 'ask_company'
    elif current_state == 'ask_company':
        session['temp_experience']['company'] = user_message
        next_state_name = 'ask_exp_dates'
    elif current_state == 'ask_exp_dates':
        session['temp_experience']['dates'] = user_message
        next_state_name = 'ask_exp_desc'
    elif current_state == 'ask_exp_desc':
        session['temp_experience']['desc'] = user_message
        if 'experience_list' not in session['resume_data']:
            session['resume_data']['experience_list'] = []
        session['resume_data']['experience_list'].append(session['temp_experience'])
        session['temp_experience'] = {}
        next_state_name = 'ask_more_experience'
    elif current_state == 'ask_more_experience':
        if user_message.lower() == 'yes':
            next_state_name = 'ask_experience'
        else:
            next_state_name = 'ask_skills' # Move to next major section

    # --- Standard Flow ---
    elif current_state != 'complete':
        # Find the next state in the main flow
        current_index = [i for i, s in enumerate(CONVERSATION_FLOW) if s['state'] == current_state][0]
        next_state_name = CONVERSATION_FLOW[current_index + 1]['state']

    # --- Final State: Generate & Save ---
    if next_state_name == 'complete':
        session['state'] = 'complete'
        
        # 1. Save to Oracle DB
        resume_id = save_resume_to_oracle(session['resume_data'])
        if not resume_id:
            return jsonify({'response': 'Error: Could not save your resume to the database.'})
        
        # 2. Generate Files
        base_filename = f"resume_{resume_id}_{session['session_id'][:8]}"
        docx_file = f"{base_filename}.docx"
        pdf_file = f"{base_filename}.pdf"
        
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_file)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_file)
        
        try:
            generate_docx(session['resume_data'], docx_path)
            generate_pdf(session['resume_data'], pdf_path)
        except Exception as e:
            print(f"Error generating files: {e}")
            return jsonify({'response': 'Error: Could not generate resume files.'})

        # 3. Send Download Links
        download_links = f"""
        Your resume is ready!
        <br>
        <a href="/download/{docx_file}" target="_blank">Download .DOCX</a>
        <br>
        <a href="/download/{pdf_file}" target="_blank">Download .PDF</a>
        """
        bot_response = get_current_state_info('complete')['prompt'] + "<br>" + download_links
    
    else:
        # Continue conversation
        session['state'] = next_state_name
        bot_response = get_current_state_info(next_state_name)['prompt']

    return jsonify({'response': bot_response})

# --- Export Module ---
@app.route("/download/<filename>")
def download(filename):
    """Provides the generated files for download."""
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

if __name__ == "__main__":
    if not pool:
        print("Application cannot start. Oracle DB connection pool failed to initialize.")
    else:
        print("Oracle DB connection pool started.")
        app.run(debug=True)